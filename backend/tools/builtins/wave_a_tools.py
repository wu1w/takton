"""
Wave A 办公/多模态工具：
- doc_read / doc_write
- image_generate
- calendar (read/write ICS local)
- tts
- capability_status
"""
from __future__ import annotations

import asyncio
import base64
import json
import logging
import os
import re
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from backend.tools.base import BaseTool, ToolRiskLevel, ToolSource

logger = logging.getLogger(__name__)

_HOME = Path(os.path.expanduser("~/.takton"))
_CAL_DIR = _HOME / "calendar"
_CAL_FILE = _CAL_DIR / "events.json"
_MEDIA_DIR = _HOME / "media"
_OUT_DIR = Path("workspace") / "wave_a"


def _ensure_dirs() -> None:
    _CAL_DIR.mkdir(parents=True, exist_ok=True)
    _MEDIA_DIR.mkdir(parents=True, exist_ok=True)
    _OUT_DIR.mkdir(parents=True, exist_ok=True)


def _local_today() -> str:
    return datetime.now().astimezone().strftime("%Y-%m-%d")


def _load_events() -> list[dict[str, Any]]:
    _ensure_dirs()
    if not _CAL_FILE.exists():
        return []
    try:
        data = json.loads(_CAL_FILE.read_text(encoding="utf-8"))
        return list(data) if isinstance(data, list) else []
    except Exception:
        return []


def _save_events(events: list[dict[str, Any]]) -> None:
    _ensure_dirs()
    _CAL_FILE.write_text(json.dumps(events, ensure_ascii=False, indent=2), encoding="utf-8")
    # also write ICS
    ics_path = _CAL_DIR / "takton.ics"
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Takton//Calendar//CN",
        "CALSCALE:GREGORIAN",
    ]
    for e in events:
        if e.get("deleted"):
            continue
        uid = e.get("id") or uuid.uuid4().hex
        start = str(e.get("start") or "").replace("-", "").replace(":", "").replace(" ", "T")
        end = str(e.get("end") or start).replace("-", "").replace(":", "").replace(" ", "T")
        if len(start) == 8:
            start = start + "T090000"
        if len(end) == 8:
            end = end + "T100000"
        summary = str(e.get("title") or "Event").replace("\n", " ")
        desc = str(e.get("description") or "").replace("\n", "\\n")
        lines += [
            "BEGIN:VEVENT",
            f"UID:{uid}@takton",
            f"DTSTART:{start}",
            f"DTEND:{end}",
            f"SUMMARY:{summary}",
            f"DESCRIPTION:{desc}",
            "END:VEVENT",
        ]
    lines.append("END:VCALENDAR")
    ics_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _resolve_path(path: str) -> Path | None:
    p = (path or "").strip().strip('"')
    if not p:
        return None
    path_obj = Path(p)
    if not path_obj.is_absolute():
        # try cwd / workspace
        for base in (Path.cwd(), Path("workspace"), _MEDIA_DIR, _OUT_DIR):
            cand = (base / p).resolve()
            if cand.exists():
                return cand
        return (Path.cwd() / p).resolve()
    return path_obj


# ── capability_status ──────────────────────────────────────


class CapabilityStatusTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="capability_status",
            description=(
                "探测 Wave A/B 能力通道是否可用（文档解析、图生、TTS、日历、浏览器等）。"
                "缺依赖时给出安装提示。"
            ),
            parameters={"type": "object", "properties": {"detail": {"type": "boolean", "default": True}}},
            source=ToolSource.BUILTIN,
            risk_level=ToolRiskLevel.SAFE,
        )

    async def execute(self, **kwargs: Any) -> Any:
        import importlib.util as iu

        def has(mod: str) -> bool:
            return iu.find_spec(mod) is not None

        checks = {
            "doc_pdf": has("fitz") or has("pypdf") or has("PyPDF2"),
            "doc_docx": has("docx"),
            "doc_xlsx": has("openpyxl"),
            "tts_edge": has("edge_tts"),
            "image_pil": has("PIL"),
            "playwright": has("playwright"),
            "pywinauto": has("pywinauto"),
        }
        env = {
            "FAL_KEY": bool(os.environ.get("FAL_KEY") or os.environ.get("FAL_API_KEY")),
            "OPENAI_API_KEY": bool(os.environ.get("OPENAI_API_KEY")),
            "TAVILY_API_KEY": bool(os.environ.get("TAVILY_API_KEY") or os.environ.get("SEARCH_API_KEY")),
            "IM_WEBHOOK_URL": bool(os.environ.get("IM_WEBHOOK_URL")),
        }
        install_hints = []
        if not checks["doc_pdf"]:
            install_hints.append("pip install pymupdf")
        if not checks["doc_docx"]:
            install_hints.append("pip install python-docx")
        if not checks["doc_xlsx"]:
            install_hints.append("pip install openpyxl")
        if not checks["tts_edge"]:
            install_hints.append("pip install edge-tts")

        return json.dumps(
            {
                "modules": checks,
                "env": env,
                "calendar_store": str(_CAL_FILE),
                "media_dir": str(_MEDIA_DIR),
                "install_hints": install_hints,
                "wave_a_tools": [
                    "doc_read",
                    "doc_write",
                    "image_generate",
                    "calendar",
                    "tts",
                    "capability_status",
                ],
            },
            ensure_ascii=False,
            indent=2,
        )


# ── doc_read / doc_write ───────────────────────────────────


class DocReadTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="doc_read",
            description=(
                "读取 PDF/DOCX/XLSX/TXT/MD 文档正文。支持 path 或 url。"
                "大文件用 offset/limit 分页（按行或字符块）。"
            ),
            parameters={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "本地路径"},
                    "url": {"type": "string", "description": "远程 URL（下载后解析）"},
                    "format": {
                        "type": "string",
                        "enum": ["auto", "pdf", "docx", "xlsx", "txt", "md"],
                        "default": "auto",
                    },
                    "offset": {"type": "integer", "default": 0, "description": "起始行（0-based）"},
                    "limit": {"type": "integer", "default": 200, "description": "最多返回行数"},
                    "query": {"type": "string", "description": "可选：只返回含关键词的行"},
                },
            },
            source=ToolSource.BUILTIN,
            risk_level=ToolRiskLevel.SAFE,
        )

    async def execute(self, **kwargs: Any) -> Any:
        path = (kwargs.get("path") or "").strip()
        url = (kwargs.get("url") or "").strip()
        fmt = str(kwargs.get("format") or "auto").lower()
        offset = int(kwargs.get("offset") or 0)
        limit = max(1, min(int(kwargs.get("limit") or 200), 2000))
        query = (kwargs.get("query") or "").strip()

        local: Path | None = None
        if url:
            local = await asyncio.to_thread(self._download, url)
            if isinstance(local, str) and local.startswith("[Error]"):
                return local
        elif path:
            local = _resolve_path(path)
            if local is None or not local.exists():
                return f"[Error] file not found: {path}"
        else:
            return "[Error] path or url required"

        assert isinstance(local, Path)
        if fmt == "auto":
            fmt = self._guess(local)

        try:
            text = await asyncio.to_thread(self._extract, local, fmt)
        except Exception as e:
            return f"[Error] extract failed ({fmt}): {e}"

        lines = text.splitlines()
        if query:
            q = query.lower()
            lines = [ln for ln in lines if q in ln.lower()]
        total = len(lines)
        chunk = lines[offset : offset + limit]
        body = "\n".join(chunk)
        meta = {
            "path": str(local.resolve()),
            "format": fmt,
            "total_lines": total,
            "offset": offset,
            "returned": len(chunk),
            "chars": len(body),
        }
        return body + "\n\n---\n" + json.dumps(meta, ensure_ascii=False)

    def _guess(self, p: Path) -> str:
        ext = p.suffix.lower()
        return {
            ".pdf": "pdf",
            ".docx": "docx",
            ".xlsx": "xlsx",
            ".xls": "xlsx",
            ".md": "md",
            ".txt": "txt",
            ".csv": "txt",
            ".json": "txt",
            ".log": "txt",
        }.get(ext, "txt")

    def _download(self, url: str) -> Path | str:
        _ensure_dirs()
        try:
            req = Request(url, headers={"User-Agent": "TaktonDocRead/1.0"})
            with urlopen(req, timeout=60) as resp:
                data = resp.read()
                ctype = resp.headers.get("Content-Type", "")
            name = Path(urlparse(url).path).name or f"dl_{uuid.uuid4().hex[:8]}"
            if "." not in name:
                if "pdf" in ctype:
                    name += ".pdf"
                elif "word" in ctype:
                    name += ".docx"
                else:
                    name += ".bin"
            dest = _MEDIA_DIR / name
            dest.write_bytes(data)
            return dest
        except Exception as e:
            return f"[Error] download failed: {e}"

    def _extract(self, path: Path, fmt: str) -> str:
        if fmt in ("txt", "md"):
            return path.read_text(encoding="utf-8", errors="replace")
        if fmt == "pdf":
            try:
                import fitz  # pymupdf

                doc = fitz.open(path)
                parts = []
                for i, page in enumerate(doc):
                    parts.append(f"--- page {i+1} ---\n{page.get_text()}")
                doc.close()
                return "\n".join(parts)
            except ImportError:
                try:
                    from pypdf import PdfReader

                    r = PdfReader(str(path))
                    return "\n".join((p.extract_text() or "") for p in r.pages)
                except ImportError as e:
                    raise RuntimeError("install pymupdf or pypdf") from e
        if fmt == "docx":
            try:
                from docx import Document
            except ImportError as e:
                raise RuntimeError("install python-docx") from e
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text]
            for ti, table in enumerate(doc.tables):
                parts.append(f"--- table {ti+1} ---")
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(parts)
        if fmt == "xlsx":
            try:
                import openpyxl
            except ImportError as e:
                raise RuntimeError("install openpyxl") from e
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"--- sheet: {sheet.title} ---")
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i > 5000:
                        parts.append("...[truncated rows]")
                        break
                    parts.append("\t".join("" if c is None else str(c) for c in row))
            wb.close()
            return "\n".join(parts)
        return path.read_text(encoding="utf-8", errors="replace")


class DocWriteTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="doc_write",
            description=(
                "写入文档。format=md|txt|docx。返回绝对路径。"
                "简报/说明用 md；需 Word 用 docx。"
            ),
            parameters={
                "type": "object",
                "properties": {
                    "content": {"type": "string"},
                    "filename": {"type": "string"},
                    "format": {"type": "string", "enum": ["md", "txt", "docx"], "default": "md"},
                    "title": {"type": "string"},
                },
                "required": ["content"],
            },
            source=ToolSource.BUILTIN,
            risk_level=ToolRiskLevel.MEDIUM,
        )

    async def execute(self, **kwargs: Any) -> Any:
        content = str(kwargs.get("content") or "")
        if not content.strip():
            return "[Error] content required"
        fmt = str(kwargs.get("format") or "md").lower()
        title = (kwargs.get("title") or "Takton Document").strip()
        name = (kwargs.get("filename") or "").strip()
        _ensure_dirs()
        if not name:
            name = f"doc_{uuid.uuid4().hex[:8]}.{fmt if fmt != 'md' else 'md'}"
        if fmt == "docx" and not name.endswith(".docx"):
            name += ".docx"
        elif fmt == "md" and not name.endswith((".md", ".markdown")):
            name += ".md"
        elif fmt == "txt" and not name.endswith(".txt"):
            name += ".txt"

        dest = (_OUT_DIR / name).resolve()
        dest.parent.mkdir(parents=True, exist_ok=True)

        if fmt == "docx":
            try:
                from docx import Document
                from docx.shared import Pt
            except ImportError:
                # fallback md
                dest = dest.with_suffix(".md")
                dest.write_text(f"# {title}\n\n{content}", encoding="utf-8")
                return f"[warn] python-docx missing; wrote markdown instead\nOK {dest}"
            doc = Document()
            doc.add_heading(title, level=1)
            for para in content.split("\n"):
                p = doc.add_paragraph(para)
                for run in p.runs:
                    run.font.size = Pt(11)
            doc.save(str(dest))
            return f"OK {dest}"

        text = content if fmt == "txt" else (f"# {title}\n\n{content}" if not content.lstrip().startswith("#") else content)
        dest.write_text(text, encoding="utf-8")
        return f"OK {dest}"


# ── image_generate ─────────────────────────────────────────


class ImageGenerateTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="image_generate",
            description=(
                "文生图。优先 FAL_KEY/FAL_API_KEY；否则 OpenAI images；"
                "都不可用时生成 Pillow 占位图并标明。返回本地 png 路径。"
            ),
            parameters={
                "type": "object",
                "properties": {
                    "prompt": {"type": "string"},
                    "filename": {"type": "string"},
                    "width": {"type": "integer", "default": 1024},
                    "height": {"type": "integer", "default": 1024},
                },
                "required": ["prompt"],
            },
            source=ToolSource.BUILTIN,
            risk_level=ToolRiskLevel.MEDIUM,
        )

    async def execute(self, **kwargs: Any) -> Any:
        prompt = str(kwargs.get("prompt") or "").strip()
        if not prompt:
            return "[Error] prompt required"
        _ensure_dirs()
        name = (kwargs.get("filename") or f"img_{uuid.uuid4().hex[:8]}.png").strip()
        if not name.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
            name += ".png"
        dest = (_MEDIA_DIR / name).resolve()
        w = int(kwargs.get("width") or 1024)
        h = int(kwargs.get("height") or 1024)
        w = max(256, min(w, 2048))
        h = max(256, min(h, 2048))

        # 1) FAL flux schnell
        fal = (os.environ.get("FAL_KEY") or os.environ.get("FAL_API_KEY") or "").strip()
        if fal:
            try:
                img_bytes = await asyncio.to_thread(self._fal_generate, fal, prompt, w, h)
                if img_bytes:
                    dest.write_bytes(img_bytes)
                    return f"OK provider=fal path={dest}\nprompt={prompt[:200]}"
            except Exception as e:
                logger.warning("fal image failed: %s", e)

        # 2) OpenAI
        oai = (os.environ.get("OPENAI_API_KEY") or "").strip()
        if oai:
            try:
                img_bytes = await asyncio.to_thread(self._openai_generate, oai, prompt, w, h)
                if img_bytes:
                    dest.write_bytes(img_bytes)
                    return f"OK provider=openai path={dest}\nprompt={prompt[:200]}"
            except Exception as e:
                logger.warning("openai image failed: %s", e)

        # 3) Pillow placeholder
        try:
            from PIL import Image, ImageDraw, ImageFont

            img = Image.new("RGB", (w, h), color=(36, 28, 58))
            draw = ImageDraw.Draw(img)
            draw.rectangle([20, 20, w - 20, h - 20], outline=(160, 120, 255), width=3)
            # wrap prompt
            text = f"Takton placeholder\n\n{prompt[:400]}"
            y = 40
            for line in text.split("\n"):
                draw.text((40, y), line[:80], fill=(230, 220, 255))
                y += 28
            draw.text((40, h - 60), "Set FAL_KEY or OPENAI_API_KEY for real images", fill=(180, 170, 200))
            img.save(dest, format="PNG")
            return (
                f"⚠️ 占位图已生成（非 AI 真图） path={dest}\n"
                f"prompt={prompt[:200]}\n"
                "原因：未配置文生图 API Key，无法调用 FAL / OpenAI Images。\n"
                "请老板配置其一后重试 image_generate：\n"
                "  • 环境变量 FAL_KEY 或 FAL_API_KEY（推荐 fal.ai flux）\n"
                "  • 或 OPENAI_API_KEY（DALL·E）\n"
                "  • Windows 可在系统环境 / Takton .env 写入后重启后端\n"
                "也可用 capability_status 查看 image 通道是否就绪。"
            )
        except Exception as e:
            return f"[Error] image_generate failed: {e}"

    def _fal_generate(self, key: str, prompt: str, w: int, h: int) -> bytes | None:
        # fal queue API - flux/schnell
        payload = json.dumps(
            {
                "prompt": prompt,
                "image_size": {"width": w, "height": h},
                "num_images": 1,
            }
        ).encode()
        req = Request(
            "https://fal.run/fal-ai/flux/schnell",
            data=payload,
            headers={
                "Authorization": f"Key {key}",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        with urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode())
        images = data.get("images") or []
        if not images:
            return None
        url = images[0].get("url")
        if not url:
            return None
        with urlopen(url, timeout=60) as r2:
            return r2.read()

    def _openai_generate(self, key: str, prompt: str, w: int, h: int) -> bytes | None:
        # map size
        size = "1024x1024"
        if w >= 1500 or h >= 1500:
            size = "1024x1024"
        payload = json.dumps(
            {
                "model": "dall-e-3",
                "prompt": prompt[:4000],
                "n": 1,
                "size": size,
                "response_format": "b64_json",
            }
        ).encode()
        req = Request(
            "https://api.openai.com/v1/images/generations",
            data=payload,
            headers={
                "Authorization": f"Bearer {key}",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        with urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode())
        items = data.get("data") or []
        if not items:
            return None
        b64 = items[0].get("b64_json")
        if b64:
            return base64.b64decode(b64)
        url = items[0].get("url")
        if url:
            with urlopen(url, timeout=60) as r2:
                return r2.read()
        return None


# ── calendar ───────────────────────────────────────────────


class CalendarTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="calendar",
            description=(
                "本地 ICS 日历。action=list|create|update|delete|export。"
                "create 需 title + start(YYYY-MM-DD or ISO)。数据在 ~/.takton/calendar/"
            ),
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["list", "create", "update", "delete", "export"],
                        "default": "list",
                    },
                    "title": {"type": "string"},
                    "description": {"type": "string"},
                    "start": {"type": "string", "description": "YYYY-MM-DD or YYYY-MM-DDTHH:MM"},
                    "end": {"type": "string"},
                    "date": {"type": "string", "description": "list 过滤起始日"},
                    "days": {"type": "integer", "default": 7},
                    "event_id": {"type": "string"},
                },
            },
            source=ToolSource.BUILTIN,
            risk_level=ToolRiskLevel.MEDIUM,
        )

    async def execute(self, **kwargs: Any) -> Any:
        action = str(kwargs.get("action") or "list").lower()
        events = _load_events()

        if action == "export":
            _save_events(events)
            return f"OK ics={(_CAL_DIR / 'takton.ics').resolve()} json={_CAL_FILE.resolve()}"

        if action == "list":
            start_day = (kwargs.get("date") or _local_today()).strip()
            days = max(1, min(int(kwargs.get("days") or 7), 90))
            try:
                d0 = datetime.strptime(start_day[:10], "%Y-%m-%d").date()
            except ValueError:
                d0 = datetime.now().date()
            d1 = d0 + timedelta(days=days)
            rows = []
            for e in events:
                if e.get("deleted"):
                    continue
                s = str(e.get("start") or "")[:10]
                try:
                    ed = datetime.strptime(s, "%Y-%m-%d").date()
                except ValueError:
                    continue
                if d0 <= ed < d1:
                    rows.append(e)
            rows.sort(key=lambda x: str(x.get("start") or ""))
            if not rows:
                return f"No events from {d0} for {days} day(s). Store: {_CAL_FILE}"
            lines = [f"# Calendar {d0} +{days}d ({len(rows)} events)"]
            for e in rows:
                lines.append(
                    f"- [{e.get('id')}] {e.get('start')} ~ {e.get('end')}: {e.get('title')}"
                    + (f" — {e.get('description')}" if e.get("description") else "")
                )
            return "\n".join(lines)

        if action == "create":
            title = (kwargs.get("title") or "").strip()
            start = (kwargs.get("start") or "").strip()
            if not title or not start:
                return "[Error] title and start required"
            end = (kwargs.get("end") or "").strip()
            if not end:
                # +1h if time present else same day
                if "T" in start:
                    try:
                        dt = datetime.fromisoformat(start)
                        end = (dt + timedelta(hours=1)).isoformat(timespec="minutes")
                    except Exception:
                        end = start
                else:
                    end = start
            eid = "evt_" + uuid.uuid4().hex[:10]
            ev = {
                "id": eid,
                "title": title,
                "description": str(kwargs.get("description") or ""),
                "start": start,
                "end": end,
                "created_at": datetime.now(timezone.utc).isoformat(),
            }
            events.append(ev)
            _save_events(events)
            return f"OK created {eid}\n{json.dumps(ev, ensure_ascii=False, indent=2)}\nics={_CAL_DIR / 'takton.ics'}"

        if action == "update":
            eid = (kwargs.get("event_id") or "").strip()
            if not eid:
                return "[Error] event_id required"
            found = next((e for e in events if e.get("id") == eid and not e.get("deleted")), None)
            if not found:
                return f"[Error] event not found: {eid}"
            for k in ("title", "description", "start", "end"):
                if kwargs.get(k) is not None and str(kwargs.get(k)).strip() != "":
                    found[k] = str(kwargs.get(k)).strip()
            _save_events(events)
            return f"OK updated {eid}\n{json.dumps(found, ensure_ascii=False, indent=2)}"

        if action == "delete":
            eid = (kwargs.get("event_id") or "").strip()
            if not eid:
                return "[Error] event_id required"
            found = next((e for e in events if e.get("id") == eid and not e.get("deleted")), None)
            if not found:
                return f"[Error] event not found: {eid}"
            found["deleted"] = True
            _save_events(events)
            return f"OK deleted {eid}"

        return f"[Error] unknown action={action}"


# ── tts ────────────────────────────────────────────────────


class TtsTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="tts",
            description=(
                "文本转语音。默认 Edge-TTS（免 key）。"
                "voice 例: zh-CN-XiaoxiaoNeural / en-US-JennyNeural。"
                "返回 mp3 绝对路径。"
            ),
            parameters={
                "type": "object",
                "properties": {
                    "text": {"type": "string"},
                    "voice": {"type": "string", "default": "zh-CN-XiaoxiaoNeural"},
                    "filename": {"type": "string"},
                },
                "required": ["text"],
            },
            source=ToolSource.BUILTIN,
            risk_level=ToolRiskLevel.LOW,
        )

    async def execute(self, **kwargs: Any) -> Any:
        text = str(kwargs.get("text") or "").strip()
        if not text:
            return "[Error] text required"
        if len(text) > 5000:
            text = text[:5000] + "…"
        voice = (kwargs.get("voice") or "zh-CN-XiaoxiaoNeural").strip()
        _ensure_dirs()
        name = (kwargs.get("filename") or f"tts_{uuid.uuid4().hex[:8]}.mp3").strip()
        if not name.lower().endswith(".mp3"):
            name += ".mp3"
        dest = (_MEDIA_DIR / name).resolve()

        try:
            import edge_tts
        except ImportError:
            return "[Error] edge-tts not installed. Run: pip install edge-tts"

        try:
            communicate = edge_tts.Communicate(text, voice)
            await communicate.save(str(dest))
            return f"OK path={dest}\nvoice={voice}\nchars={len(text)}"
        except Exception as e:
            return f"[Error] tts failed: {e}"


WAVE_A_TOOL_CLASSES = [
    CapabilityStatusTool,
    DocReadTool,
    DocWriteTool,
    ImageGenerateTool,
    CalendarTool,
    TtsTool,
]
