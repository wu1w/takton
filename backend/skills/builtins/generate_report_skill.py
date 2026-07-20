"""
报告生成 Skill
根据用户提供的主题和要求，生成结构化报告（优先 Word .docx，并保留 .md）
"""

import logging
import os
import uuid
from datetime import datetime, timezone

from backend.skills.base import BaseSkill

logger = logging.getLogger(__name__)

REPORT_OUTPUT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), "..", "..", "..", "uploads", "reports")
)

SYSTEM_PROMPT = """你是一位专业的报告撰写专家。你的任务是根据用户提供的主题和要求，生成一份结构清晰、内容详实的专业报告。

报告必须包含以下结构：
1. 执行摘要（Executive Summary）
2. 背景介绍（Background）
3. 主要内容（Main Content）- 分章节展开
4. 数据分析/案例研究（如适用）
5. 结论与建议（Conclusion & Recommendations）
6. 附录（Appendix，如适用）

格式要求：
1. 使用Markdown格式输出
2. 标题层级清晰（# ## ###）
3. 适当使用表格、列表增强可读性
4. 内容专业、数据准确、逻辑严密
5. 报告长度根据主题复杂度调整，一般不少于2000字
"""


class GenerateReportSkill(BaseSkill):
    """生成专业报告"""

    name = "generate_report"
    description = (
        "根据主题生成专业报告（市场分析/技术/调研等）。"
        "输出优先 Word(.docx)，并保留 Markdown 副本。"
    )
    parameters = {
        "type": "object",
        "properties": {
            "topic": {"type": "string", "description": "报告主题"},
            "report_type": {
                "type": "string",
                "description": "报告类型",
                "enum": [
                    "market_analysis",
                    "technical",
                    "research",
                    "business",
                    "summary",
                ],
                "default": "summary",
            },
            "language": {"type": "string", "description": "报告语言", "default": "zh"},
            "sections": {
                "type": "string",
                "description": "可选的章节要求，用逗号分隔",
                "default": "",
            },
        },
        "required": ["topic"],
    }

    async def execute(
        self,
        topic: str,
        report_type: str = "summary",
        language: str = "zh",
        sections: str = "",
        **kwargs,
    ) -> str:
        try:
            from backend.services.llm import LLMServiceFactory

            type_names = {
                "market_analysis": "市场分析报告",
                "technical": "技术报告",
                "research": "调研报告",
                "business": "商业计划书",
                "summary": "总结报告",
            }
            type_name = type_names.get(report_type, "专业报告")

            prompt = (
                f"请生成一份{type_name}。\n\n主题：{topic}\n"
                f"语言：{'中文' if language == 'zh' else 'English'}\n"
            )
            if sections:
                prompt += f"必须包含的章节：{sections}\n"
            prompt += f"\n请严格按照以下要求撰写报告：\n{SYSTEM_PROMPT}\n"
            prompt += "\n请直接输出Markdown格式的报告内容，不要添加额外的解释文字。"

            llm_service = LLMServiceFactory.get_service()
            response = ""
            async for chunk in llm_service.chat(
                [{"role": "user", "content": prompt}], stream=False
            ):
                response += chunk.delta or ""
                if chunk.finish_reason:
                    break

            if not response.strip():
                return "[Error] 报告生成失败：LLM 未返回任何内容"

            report_md = self._extract_markdown(response)
            if not report_md.strip():
                return "[Error] 报告生成失败：解析后的报告内容为空"

            file_path = self._save_report_file(report_md, topic, report_type)
            file_name = os.path.basename(file_path)
            summary = self._extract_summary(report_md)
            fmt = (
                "Word .docx"
                if str(file_path).endswith(".docx")
                else "Markdown .md（安装 python-docx 可出 Word）"
            )

            return (
                f"[Success] 报告生成完成！\n"
                f"类型：{type_name}\n"
                f"主题：{topic}\n"
                f"摘要：{summary}\n"
                f"文件路径：{file_path}\n"
                f"格式：{fmt}\n"
                f"下载：/uploads/reports/{file_name}"
            )

        except Exception as e:
            logger.error(f"Report generation failed: {e}")
            return f"[Error] 报告生成失败: {e}"

    def _extract_markdown(self, text: str) -> str:
        import re

        code_block = re.search(r"```(?:markdown)?\s*(.*?)\s*```", text, re.DOTALL)
        if code_block:
            return code_block.group(1).strip()
        return text.strip()

    def _save_report_file(self, report_md: str, topic: str, report_type: str) -> str:
        """保存报告文件"""
        # 添加报告头信息
        header = f"""# {topic}

> **报告类型**：{report_type}  
> **生成时间**：{datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M:%S %Z')}  
> **生成工具**：Takton Agent

---

"""
        full_content = header + report_md

        os.makedirs(REPORT_OUTPUT_DIR, exist_ok=True)
        stem = uuid.uuid4().hex
        md_path = os.path.join(REPORT_OUTPUT_DIR, f"{stem}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(full_content)

        docx_path = os.path.join(REPORT_OUTPUT_DIR, f"{stem}.docx")
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            doc.add_heading(topic or "报告", level=0)
            doc.add_paragraph(
                f"报告类型：{report_type}  |  生成时间："
                f"{datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')}"
            )
            for line in report_md.splitlines():
                s = line.strip()
                if not s:
                    doc.add_paragraph("")
                    continue
                if s.startswith("### "):
                    doc.add_heading(s[4:], level=3)
                elif s.startswith("## "):
                    doc.add_heading(s[3:], level=2)
                elif s.startswith("# "):
                    doc.add_heading(s[2:], level=1)
                elif s.startswith(("- ", "* ")):
                    p = doc.add_paragraph(s[2:], style="List Bullet")
                    for run in p.runs:
                        run.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(s)
                    for run in p.runs:
                        run.font.size = Pt(11)
            doc.save(docx_path)
            return docx_path
        except Exception as e:
            logger.warning("docx export failed, keep md: %s", e)
            return md_path

    def _extract_summary(self, report_md: str, max_length: int = 200) -> str:
        lines = report_md.split("\n")
        for line in lines:
            stripped = line.strip()
            if stripped and not stripped.startswith("#") and not stripped.startswith("---"):
                if len(stripped) > max_length:
                    return stripped[:max_length] + "..."
                return stripped
        return "报告已生成"
